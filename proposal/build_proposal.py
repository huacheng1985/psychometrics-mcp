from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "NCME_2027_Psychometrics_MCP_Proposal.docx"

TITLE = "Psychometrics MCP: Measurement-Aware Tools for AI-Assisted Analysis"
ABSTRACT = (
    "Psychometrics MCP is an open-source, local-first server that lets AI assistants invoke "
    "constrained psychometric analyses with explicit diagnostics, provenance, and interpretation "
    "limits. This demonstration shows data inspection, classical item analysis, and Rasch "
    "estimation, then teaches attendees how measurement-aware tool design can improve "
    "reproducibility without surrendering expert judgment or sensitive data."
)
SUMMARY_PARAGRAPHS = [
    (
        "AI assistants can generate plausible statistical advice, but general-purpose tool access "
        "can also enable inappropriate models, hidden data handling, irreproducible commands, and "
        "interpretations that outrun evidence. Psychometrics MCP helps measurement researchers, "
        "practitioners, and instructors conduct reproducible, measurement-aware analysis by giving "
        "AI hosts a constrained set of validated tools with strict inputs, explicit diagnostics, "
        "provenance, and interpretation limits. The server is free, open-source, and local-first; "
        "its core analyses do not require the project to receive users' response data or pay for "
        "their language-model use."
    ),
    (
        "The working prototype implements five Model Context Protocol tools. It checks computation "
        "capabilities; audits response dimensions, missingness, category use, ranges, and zero "
        "variance; calculates classical item summaries, item-rest correlations, coefficient alpha, "
        "and standard error of measurement; generates purpose- and design-aware analysis plans; "
        "and "
        "fits a dichotomous Rasch model through a fixed eRm::RM adapter using conditional maximum "
        "likelihood. The server rejects unknown input fields and invalid response codes, returns "
        "sample flow and software versions, and never accepts arbitrary R or shell commands. Rasch "
        "is the first numerical validation slice, not the project boundary."
    ),
    (
        "During the eBoard demonstration, attendees will follow a synthetic item-response dataset "
        "through three views. First, a schema and data audit will show how the server identifies "
        "missingness, invalid categories, and unstable conditions before estimation. Second, "
        "classical and Rasch outputs will illustrate the difference between obtaining a number and "
        "obtaining an auditable result with assumptions, exclusions, diagnostics, and "
        "interpretation boundaries. Third, a live tool call from an AI host will show that the "
        "host "
        "can request an analysis while the fixed analytical engine, rather than model-generated "
        "code, controls what executes. A local Docker deployment will demonstrate that protected "
        "response data can remain "
        "on the user's computer."
    ),
    (
        "After the demonstration, attendees will be able to: distinguish constrained analytical "
        "tools from unrestricted code execution; identify the provenance and diagnostic fields "
        "needed to audit AI-assisted psychometric results; run the open-source server locally; and "
        "adapt the design pattern to additional methods without treating model fit, reliability, "
        "or "
        "prediction as automatic validity evidence. Attendees will also receive a development map "
        "covering descriptive statistics, correlation, regression, factor models, broader "
        "Rasch/IRT, DIF and invariance, linking and equating, CAT/MST, rater and generalizability "
        "models, measurement-aware machine learning, and reproducible reporting."
    ),
    (
        "The innovation addresses the 2027 meeting theme by connecting measurement expertise with "
        "AI infrastructure while preserving human oversight. The repository already includes MCP "
        "client regression tests, a real R/eRm numerical integration test, container deployment, "
        "privacy boundaries, and continuous-integration workflows. Before the Annual Meeting, "
        "independently verified reference datasets, additional numerical benchmarks, user "
        "documentation, and a reusable demonstration dataset will be released. The result will be "
        "a "
        "practical community resource rather than a commercial product or a claim that AI can "
        "replace psychometric judgment."
    ),
]


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w]+(?:[-'][\w]+)*\b", text))


def set_font(run, size: float, color: str = "000000", bold: bool = False, italic: bool = False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa: list[int]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def shade_cell(cell, fill: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def mark_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, 9, "6B7280")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def add_heading(doc, text: str, level: int = 1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_body(doc, text: str):
    paragraph = doc.add_paragraph(text, style="Normal")
    paragraph.paragraph_format.keep_together = True
    return paragraph


def add_bullet(doc, text: str):
    paragraph = doc.add_paragraph(text, style="List Bullet")
    paragraph.paragraph_format.keep_together = True
    return paragraph


def add_number(doc, text: str):
    paragraph = doc.add_paragraph(text, style="List Number")
    paragraph.paragraph_format.keep_together = True
    return paragraph


def build():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, "2E74B5"),
        ("Heading 2", 13, 12, 6, "2E74B5"),
        ("Heading 3", 12, 8, 4, "1F4D78"),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("NCME 2027 INNOVATION DEMONSTRATION")
    set_font(run, 9, "6B7280", bold=True)
    add_page_number(section.footer.paragraphs[0])

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(10)
    set_font(kicker.add_run("PROPOSAL FOR BLIND REVIEW"), 10, "2E74B5", bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    set_font(title.add_run(TITLE), 24, "0B2545", bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    set_font(
        subtitle.add_run("Free, local-first infrastructure for auditable AI-assisted measurement"),
        12.5,
        "4B5563",
        italic=True,
    )

    metadata = document.add_table(rows=3, cols=2)
    set_table_widths(metadata, [4680, 4680])
    mark_table_header(metadata.rows[0])
    values = [
        ("Submission category", "Innovation Demonstration"),
        ("Format", "60-minute individual eBoard"),
        ("Modality", "In-person only"),
        ("Review", "Blinded"),
        ("Meeting", "Toronto | April 14–17, 2027"),
        ("Deadline", "September 13, 2026 | 11:59 PM PDT"),
    ]
    for cell, (label, value) in zip(
        [cell for row in metadata.rows for cell in row.cells], values, strict=True
    ):
        shade_cell(cell, "F4F6F9")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label + "\n"), 8.5, "6B7280", bold=True)
        set_font(p.add_run(value), 10.5, "111827")

    add_heading(document, "Abstract", 1)
    label = document.add_paragraph()
    label.paragraph_format.space_after = Pt(4)
    set_font(
        label.add_run(f"{word_count(ABSTRACT)} words | 50-word maximum"), 9, "6B7280", italic=True
    )
    add_body(document, ABSTRACT)

    notice = document.add_paragraph()
    notice.paragraph_format.space_before = Pt(10)
    notice.paragraph_format.space_after = Pt(0)
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        notice.add_run(
            "No author names or identifying institutional information appear in this proposal."
        ),
        9.5,
        "7A5A00",
        bold=True,
    )

    document.add_page_break()
    add_heading(document, "Demonstration Summary", 1)
    summary = " ".join(SUMMARY_PARAGRAPHS)
    label = document.add_paragraph()
    label.paragraph_format.space_after = Pt(4)
    set_font(
        label.add_run(f"{word_count(summary)} words | 500-word maximum"), 9, "6B7280", italic=True
    )
    for paragraph in SUMMARY_PARAGRAPHS:
        add_body(document, paragraph)

    add_heading(document, "Software Requirements", 1)
    for item in (
        (
            "Presenter: laptop with Docker, a local MCP-compatible AI host, internet access for "
            "fallback only, and an eBoard connection."
        ),
        (
            "Attendees: no software installation, account, or API key is required to follow the "
            "demonstration."
        ),
        (
            "Distribution: public source repository, container instructions, synthetic example "
            "data, and verification notes."
        ),
        (
            "Privacy: demonstration data will be synthetic or openly licensed; no student or "
            "examinee records will be uploaded."
        ),
    ):
        add_bullet(document, item)

    add_heading(document, "Development Commitments Before April 2027", 1)
    for item in (
        "Publish the open-source repository and tagged container release.",
        "Add independently verified reference datasets and numerical tolerances.",
        "Document local setup for major MCP hosts and a no-sensitive-data demonstration path.",
        "Complete a security and dependency review and publish the data-handling boundary.",
        "Prepare an accessible eBoard walkthrough and a downloadable quick-start guide.",
    ):
        add_number(document, item)

    add_heading(document, "References", 1)
    for reference in (
        (
            "Mair, P., & Hatzinger, R. (2007). Extended Rasch modeling: The eRm package for the "
            "application of IRT models in R. Journal of Statistical Software, 20(9), 1–20."
        ),
        "Model Context Protocol. (2026). Model Context Protocol specification. https://modelcontextprotocol.io/",
        (
            "National Council on Measurement in Education. (2026). 2027 NCME Annual Meeting call "
            "for proposals. "
            "https://ncme.org/wp-content/uploads/2026/08/NCME_cfp_2027_final_r1.pdf"
        ),
    ):
        p = add_body(document, reference)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)

    document.core_properties.title = TITLE
    document.core_properties.subject = "NCME 2027 Innovation Demonstration proposal"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = "Prepared for human review before submission."
    document.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print(f"Title words: {word_count(TITLE)}")
    print(f"Abstract words: {word_count(ABSTRACT)}")
    print(f"Summary words: {word_count(summary)}")


if __name__ == "__main__":
    build()
